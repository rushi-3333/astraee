"""
Build ASTRAE Final Project Documentation — single DOCX with consistent formatting.
Times New Roman: Main heading 18pt, Sub heading 14pt, Body 12pt, Page numbers in footer.
"""
from __future__ import annotations

import os
import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor

BASE = Path(__file__).resolve().parent
OUTPUT = BASE / "ASTRAE_Final_Project_Documentation.docx"

FONT = "Times New Roman"
SIZE_MAIN = 18
SIZE_SUB = 14
SIZE_BODY = 12

# Section order: (display title, source filename or None for generated)
SECTIONS = [
    ("Title Page", None),
    ("Abstract", "Abstract.docx"),
    ("Introduction", "Introduction.docx"),
    ("Literature Survey", "Literature Survey.docx"),
    ("Existing System and Proposed System", "Existing System & Proposed System.docx"),
    ("System Study", "System Study.docx"),
    ("System Specification", "System Specification.docx"),
    ("Software Environment", "Software Environment django.docx"),
    ("Methodology", "methodology.docx"),
    ("Dataset Collection", "Dataset collection.docx"),
    ("Modules", "Modules.docx"),
    ("Implemented Features (Current Build)", None),
    ("System Design", "system design.docx"),
    ("Input and Output Design", "Input and Output Design.docx"),
    ("Application Flow", "Astrae-Flow.docx"),
    ("Equations", "equations.docx"),
    ("System Testing", "System testing.docx"),
    ("Future Enhancements", "Future Enhancements.docx"),
    ("Conclusion", "Conclusion.docx"),
    ("References", "References.docx"),
]

MAIN_HEADING_RE = re.compile(
    r"^(abstract|introduction|conclusion|literature survey|methodology|"
    r"existing system|proposed system|system study|system specification|"
    r"input design|output design|references|future enhancements|"
    r"data collection|equations|system testing|modules|software environment)$",
    re.I,
)
SUB_HEADING_RE = re.compile(r"^(\d+\.|\d+\.\d+\.?|[A-Z]\.)\s+\S")


def set_run_font(run, size=SIZE_BODY, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def classify_paragraph(text: str, style_name: str) -> str:
    t = text.strip()
    if not t:
        return "empty"
    if style_name and style_name.startswith("Heading 1"):
        return "main"
    if style_name and style_name.startswith("Heading 2"):
        return "sub"
    if style_name and style_name.startswith("Heading 3"):
        return "sub"
    if MAIN_HEADING_RE.match(t) and len(t) < 80:
        return "main"
    if SUB_HEADING_RE.match(t) and len(t) < 120:
        return "sub"
    if t.isupper() and len(t) < 60:
        return "sub"
    return "body"


def add_formatted_paragraph(doc, text, kind="body", alignment=None):
    if not text or not text.strip():
        return
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text.strip())
    if kind == "main":
        set_run_font(run, SIZE_MAIN, bold=True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    elif kind == "sub":
        set_run_font(run, SIZE_SUB, bold=True)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    else:
        set_run_font(run, SIZE_BODY)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
    return p


def add_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        set_run_font(run, 10)

        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_begin)

        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        run._r.append(instr)

        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        run._r.append(fld_sep)

        run2 = p.add_run("1")
        set_run_font(run2, 10)

        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run2._r.append(fld_end)


def copy_paragraphs_from_source(doc, source_path: Path, max_paras: int | None = None, skip_title: str | None = None):
    if not source_path.exists():
        add_formatted_paragraph(doc, f"[Source file not found: {source_path.name}]", "body")
        return
    src = Document(source_path)
    count = 0
    skipped_title = False
    skip_norm = (skip_title or "").strip().lower()
    for para in src.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if not skipped_title and skip_norm:
            if text.strip().lower() == skip_norm or skip_norm in text.strip().lower()[: len(skip_norm) + 5]:
                skipped_title = True
                continue
        kind = classify_paragraph(text, para.style.name if para.style else "")
        # Avoid duplicate main headings from source when we already added section title
        if kind == "main" and skipped_title and count == 0:
            continue
        add_formatted_paragraph(doc, text, kind)
        count += 1
        if max_paras and count >= max_paras:
            add_formatted_paragraph(
                doc,
                "[Additional flow details available in the application and README.]",
                "body",
            )
            break

    # Copy images from source (system design etc.)
    for rel in src.part.rels.values():
        if "image" in rel.reltype:
            try:
                img_part = rel.target_part
                blob = img_part.blob
                # Insert after last paragraph — skip if too many
                pass  # images handled separately for system design
            except Exception:
                pass


def copy_images_from_source(doc, source_path: Path, max_images: int = 8):
    if not source_path.exists():
        return
    src = Document(source_path)
    img_count = 0
    for rel in src.part.rels.values():
        if "image" not in rel.reltype:
            continue
        if img_count >= max_images:
            break
        try:
            img_part = rel.target_part
            blob = img_part.blob
            ext = img_part.content_type.split("/")[-1]
            if ext == "jpeg":
                ext = "jpg"
            tmp = BASE / f"_tmp_img_{img_count}.{ext}"
            tmp.write_bytes(blob)
            doc.add_picture(str(tmp), width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            tmp.unlink(missing_ok=True)
            img_count += 1
        except Exception:
            continue


def add_title_page(doc):
    for _ in range(6):
        doc.add_paragraph()
    add_formatted_paragraph(
        doc,
        "ASTRAE",
        "main",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_formatted_paragraph(
        doc,
        "Unified Multi-Service Comparison and Rewards Marketplace",
        "sub",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_formatted_paragraph(
        doc,
        "Adaptive Semantic Retrieval and AI Recommendation Engine",
        "body",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    doc.add_paragraph()
    add_formatted_paragraph(
        doc,
        "Major Project Report",
        "sub",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    doc.add_paragraph()
    add_formatted_paragraph(
        doc,
        "Submitted in partial fulfilment of the requirements for the award of",
        "body",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_formatted_paragraph(
        doc,
        "Bachelor of Technology / Master of Technology",
        "body",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    doc.add_paragraph()
    add_formatted_paragraph(
        doc,
        "Department of Computer Science and Engineering",
        "body",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_formatted_paragraph(
        doc,
        "Academic Year 2025–2026",
        "body",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    doc.add_page_break()


def add_implemented_features(doc):
    add_formatted_paragraph(doc, "Implemented Features (Current Build)", "main")
    add_formatted_paragraph(
        doc,
        "The following sections describe the fully implemented ASTRAE application as of "
        "August 2026. Earlier draft documents may not include all modules listed here.",
        "body",
    )
    features = [
        ("Unified Search and Compare", [
            "Single search bar with intelligent category detection (ride, food, grocery, shopping, fashion, beauty, medicine).",
            "Parallel mock provider search with optional ChromaDB and CatBoost ML pipeline.",
            "ASTRAE Score blending price, discount, cashback, rating, delivery, and coupon factors.",
            "AIRE reinforcement-learning ranking overlay when model artifacts are available.",
            "Transparent recommendation explanations (Why ASTRAE recommends this).",
            "Graceful partial-failure messaging when a provider is unavailable.",
        ]),
        ("Events and Offers", [
            "Platform events: festival sales, flash sales, live offers, platform-specific promotions.",
            "Filters: platform, category, upcoming, live now, ending soon.",
            "Event detail page with integrated date and time-slot booking.",
            "Demo/sample data clearly labeled (is_demo=True).",
        ]),
        ("Booking and Orders", [
            "Full booking flow: date, time slot, pickup/delivery address, quantity before confirmation.",
            "Server-side price validation — client-submitted prices are never trusted.",
            "Orders stored with schedule, savings, cashback, and linked rewards.",
            "Reschedule from Orders page.",
        ]),
        ("Rewards and Savings", [
            "Configurable reward rules (order completed, coupon sold, daily login, referral).",
            "Wallet with available, lifetime, used, and pending points.",
            "My Savings dashboard with real data from order astrae_savings fields.",
            "Monthly and category breakdown charts (Chart.js).",
        ]),
        ("Coupon Marketplace", [
            "Lifecycle: verify → list → purchase → ownership transfer → redeem/expire.",
            "Atomic purchase with wallet locking and insufficient-balance protection.",
            "Demo coupon verification and DEMO-* coupon codes.",
            "Verified badge and expiry labels on marketplace listings.",
        ]),
        ("Watchlist and Price Alerts", [
            "Save items from compare results to wishlist.",
            "Create price alerts from watchlist items.",
            "Alert types: price drop, coupon expiry, deal expiry, event starting.",
            "Management command: check_price_alerts for demo notifications.",
        ]),
        ("Personalization and Notifications", [
            "Personalized home (Picked for you) vs Trending based on user activity.",
            "In-app notifications for orders, rewards, marketplace, price drops.",
            "User profile with preferences and notification settings.",
        ]),
        ("Admin Dashboard", [
            "User activation/deactivation (POST-protected).",
            "Metrics: users, searches, orders, coupons, sales, rewards, total savings.",
            "Platform health table and category/platform charts.",
            "Order, coupon, and marketplace sales logs.",
        ]),
        ("Security (Audit Hardening)", [
            "CSRF protection on all forms; login_required on user actions.",
            "Staff-only admin routes; ownership validation on user resources.",
            "Environment-based Django settings (SECRET_KEY, DEBUG, ALLOWED_HOSTS).",
            "Password change requires current password; safe redirect on wishlist add.",
        ]),
        ("Technology Stack", [
            "Backend: Django 6.1, SQLite (development).",
            "Frontend: Tailwind CSS, ASTRAE design system (astrae.css, astrae.js).",
            "ML (optional): CatBoost, ChromaDB, Sentence Transformers, AIRE ranking.",
            "Demo mode: ASTRAE_DEMO_MODE=True — simulated provider data, no live APIs.",
        ]),
    ]
    for title, bullets in features:
        add_formatted_paragraph(doc, title, "sub")
        for b in bullets:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(b)
            set_run_font(run, SIZE_BODY)


def add_updated_system_spec(doc):
    add_formatted_paragraph(doc, "System Specification", "main")
    add_formatted_paragraph(doc, "1. Minimum Hardware Requirements", "sub")
    specs_hw = [
        "Processor: Intel Core i3 / AMD Ryzen 3 or higher (64-bit).",
        "RAM: 4 GB minimum; 8 GB recommended for ML model loading.",
        "Storage: 2 GB free disk space (excluding virtual environment and datasets).",
        "Display: 1280×720 minimum; responsive design supports mobile and tablet.",
        "Network: Internet connection for initial setup and optional model downloads.",
    ]
    for s in specs_hw:
        p = doc.add_paragraph(style="List Bullet")
        set_run_font(p.add_run(s), SIZE_BODY)

    add_formatted_paragraph(doc, "2. Software and Runtime Requirements", "sub")
    specs_sw = [
        "Operating System: Windows 10/11, macOS 12+, or Linux (Ubuntu 20.04+).",
        "Python: 3.12 or higher.",
        "Web Browser: Chrome, Firefox, Edge, or Safari (latest).",
        "Framework: Django 6.1.",
        "Database: SQLite 3 (development); PostgreSQL recommended for production.",
        "Key Python packages: django, chromadb, sentence-transformers, catboost, scikit-learn, pandas, joblib.",
    ]
    for s in specs_sw:
        p = doc.add_paragraph(style="List Bullet")
        set_run_font(p.add_run(s), SIZE_BODY)

    add_formatted_paragraph(doc, "3. Deployment URLs (Development)", "sub")
    urls = [
        "Landing: http://127.0.0.1:8000/",
        "User Home: /ASTRAEUser/userhome/",
        "Compare / Search: /ASTRAEUser/usersearch/",
        "Events and Offers: /ASTRAEUser/userevents/",
        "Book Offer: /ASTRAEUser/book/",
        "Coupons, Rewards, Savings, Watchlist, Alerts, Orders, Profile — under /ASTRAEUser/",
        "Admin Dashboard: /ASTRAEAdmin/adminhome/",
    ]
    for u in urls:
        p = doc.add_paragraph(style="List Bullet")
        set_run_font(p.add_run(u), SIZE_BODY)


def add_updated_testing(doc):
    add_formatted_paragraph(doc, "System Testing", "main")
    add_formatted_paragraph(
        doc,
        "System testing verifies end-to-end integration of the ASTRAE application, "
        "including search, comparison, booking, rewards, coupon marketplace, events, "
        "watchlist, alerts, savings dashboard, and admin analytics.",
        "body",
    )
    add_formatted_paragraph(doc, "Testing Performed", "sub")
    tests = [
        ("Unit and Service Tests", "Django manage.py check — no configuration issues. Migration dry-run — schema matches models."),
        ("Search and Compare", "Ride, food, shopping, medicine categories return results. Partial provider failure shows warning banner."),
        ("Booking Flow", "GET booking form with q1/q2 params; POST creates order with server-validated pricing."),
        ("Authentication", "Register, login, logout; inactive users blocked until admin activation."),
        ("Coupon Marketplace", "List, buy, ownership transfer; insufficient balance and duplicate purchase blocked."),
        ("Security", "Admin user toggle requires POST; profile password change requires current password."),
        ("UI/UX", "Responsive layout on mobile and desktop; empty, loading, and error states on major pages."),
        ("Management Commands", "seed_demo_data, seed_events, check_price_alerts execute without errors."),
    ]
    for title, desc in tests:
        add_formatted_paragraph(doc, title, "sub")
        add_formatted_paragraph(doc, desc, "body")
    add_formatted_paragraph(doc, "Test Result Summary", "sub")
    add_formatted_paragraph(
        doc,
        "All critical user journeys (search → compare → book → reward → coupon → savings) "
        "were verified. The application is ready for academic demonstration. "
        "ML components degrade gracefully to mock providers when model artifacts are absent.",
        "body",
    )


def configure_document(doc):
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(SIZE_BODY)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def build():
    doc = Document()
    configure_document(doc)

    for section_title, source_file in SECTIONS:
        if section_title == "Title Page":
            add_title_page(doc)
            continue

        if section_title == "Implemented Features (Current Build)":
            doc.add_page_break()
            add_implemented_features(doc)
            continue

        if section_title == "System Specification":
            doc.add_page_break()
            add_updated_system_spec(doc)
            continue

        if section_title == "System Testing":
            doc.add_page_break()
            add_updated_testing(doc)
            continue

        if source_file:
            doc.add_page_break()
            add_formatted_paragraph(doc, section_title, "main")
            src_path = BASE / source_file

            if section_title == "System Design":
                copy_paragraphs_from_source(doc, src_path, max_paras=40, skip_title=section_title)
                copy_images_from_source(doc, src_path, max_images=10)
            elif section_title == "Application Flow":
                add_formatted_paragraph(
                    doc,
                    "The following summarizes the ASTRAE application flow. "
                    "The complete interface specification includes onboarding, search, "
                    "compare, booking, rewards, coupons, events, watchlist, and admin modules.",
                    "body",
                )
                copy_paragraphs_from_source(doc, src_path, max_paras=120, skip_title="ASTRAE")
            else:
                copy_paragraphs_from_source(doc, src_path, skip_title=section_title.split("(")[0].strip())

    add_page_numbers(doc)
    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")
    print(f"Size: {OUTPUT.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    build()
